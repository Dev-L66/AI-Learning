import os
import json
import time
from pathlib import Path
from dotenv import load_dotenv
from pydantic import BaseModel
from groq import Groq
from pypdf import PdfReader
from docx import Document

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

client = Groq(api_key = my_api_key)

if not my_api_key:
    raise ValueError("API KEY not found.")


role = "user"
model = "llama-3.3-70b-versatile"

job_description = """
Description
Required Academic Background

Applicants should have graduated within the past 12 months with a Bachelor's, Master's, or Doctoral degree in one of the following fields:

Engineering
Information Technology
Innovation
Data Science
Economics
Business Administration
Development Studies
Required Skills
Strong research and analytical skills
Interest in technology and development solutions
Ability to synthesize and present information effectively
Proficiency in Microsoft Excel
Experience with Power BI or similar data visualization tools
Strong attention to detail and organized work style
Excellent written communication skills
Language Requirements
Good command of English (required)
Knowledge of Arabic and/or French is an advantage
Internship Objective

The internship provides practical exposure to technology deployment, market analysis, and innovation assessment within the framework of the IsDB Technology Deployment Cooperation Program.

The intern will support technology screening, benchmarking, database development, and the preparation of analytical and decision-support materials related to development-focused technology initiatives across IsDB Member Countries.

Scope of Work

Under the supervision of the division, the intern will contribute to the implementation of the Technology Deployment Cooperation Program by supporting the identification and sourcing of innovative, cost-effective technologies that address development challenges in IsDB Member Countries.

Key responsibilities include:

Establishing technology and provider databases
Preparing technology screening summaries
Developing TECH-aligned comparative analysis matrices
Conducting market sounding analyses
Preparing presentation slides and analytical reports
Supporting technology shortlisting and decision-making materials
Detailed Responsibilities
1. Market Research and Technology Sourcing
Conduct background research on priority technology sectors.
Map global technology providers, with emphasis on providers from IsDB Member Countries.
Support preparation and dissemination of market sounding calls.
Maintain internal tracking of submissions received.
2. Technology Screening and Database Management
Assist in the initial screening of submissions against eligibility criteria, including:
Market readiness
Deployment track record
Organize Technology Appraisal Forms and supporting documentation, including:
Case studies
Certifications
Financial information
Populate and maintain structured technology databases.
3. TECH-Based Analysis and Benchmarking

Support the assessment of submitted technologies using the TECH evaluation framework:

Technology – Innovation, uniqueness, and transformational potential
Environmental & Social Impact – Sustainability and social inclusion
Comprehensive Deployment – Operational maturity, deployment approach, and implementation record
Home Adaptability – Suitability for IsDB Member Countries and capacity development potential

The intern will also assist in benchmarking technologies based on:

Technical feasibility and performance evidence
Capital (CAPEX) and operational (OPEX) costs
Affordability
Scalability and deployment requirements
Risk assessment and safeguard considerations
4. Market Analysis and Knowledge Products

Assist the division in preparing:

Market analysis reports
Technology readiness assessments
Indicative pricing models
Business model analyses
Delivery approach assessments
Trend analysis and innovation gap identification

Support the development of knowledge products, including:

Market scan summaries
Technology landscape overviews
Shortlisting dashboards
Internal briefing notes
Expected Learning Outcomes

By the end of the internship, participants will gain:

Hands-on experience in technology sourcing and market sounding
Practical knowledge of technology readiness assessment
Understanding of the TECH evaluation framework
Experience supporting development-oriented technology programs
Exposure to technology benchmarking and decision-support processes

"""
class JobDescription(BaseModel):
 role: str
 required_skills: list[str]
 preferred_skills: list[str]
 minimum_experience: float | None
 education_requirements: list[str]
 responsibilities: list[str]

jobdescription_schema = JobDescription.model_json_schema()

system_prompt = f"""
You are an expert HR assistance.
Your job us to analyze job description and extract structured info from them.

Return only valid JSON matching this schema {jobdescription_schema}

IMPORTANT:
Do NOT return the schema itself.
Do NOT return fields like "properties", "title" or "type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
Do not invent information.
"""
  
user_prompt = f"""
Analyze the job description {job_description}.
"""

message_system ={
   "role": "system",
   "content": system_prompt
}
message_user ={
   "role": "user",
   "content": user_prompt
}
response_format = {
    "type": "json_object"
}


messages = [message_system, message_user]
response = client.chat.completions.create(messages = messages, model = model, response_format = response_format)
answer = response.choices[0].message.content

raw_json = answer
job_data = json.loads(raw_json)
job = JobDescription(**job_data)

print(job.minimum_experience)
print(job.education_requirements)


class MatchResult(BaseModel):
   score: float
   details: dict

class Experience(BaseModel):
   company: str | None = None
   role: str | None = None
   duration: str | None = None
   description: str | None = None
   skills_used: list[str] = []

class Resume(BaseModel):
   name:str | None = None
   email:str | None = None
   phone: str | None = None
   total_experience_years: float | None = None
   skills: list[str] = []
   experiences: list[Experience]= []
   education: list[str] = []
   projects: list[str] = []
   certifications: list[str] = []

resume_schema = Resume.model_json_schema()

def final_score(job, resume):
   match_schema = MatchResult.model_json_schema()

   prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict
    7. Overall match (good, bad)

    Keep the response concise and easy to read.
"""
   message = {
   "role" : "user",
   "content": prompt
    }
   
   messages = [message]
   response_format = {
      "type": "json_object"
   }

   response = client.chat.completions.create(model = model, messages= messages, response_format = response_format)
   data = json.loads(response.choices[0].message.content)
   return MatchResult(**data)

def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """

    message_system = {
       "role": "system",
       "content": system_prompt
     }
    
    message_user = {
       "role": "user",
       "content" : user_prompt
    }

    messages = [message_system, message_user]
    response_format = {
       "type": "json_object"
    }
    response = client.chat.completions.create(model= model, messages= messages, response_format= response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume


def read_pdf(file_path):
  reader = PdfReader(file_path)
  text = "" 

  for page in reader.pages:
     page_text = page.extract_text()
     if page_text:
       text += page_text + "\n"
  return text

def read_docx(file_path):
   document = Document(file_path)
   text = ""
   for paragraph in document.paragraphs:
      if paragraph.text.strip():
         text  += paragraph.text + "\n"

   for table in document.tables:
      for row in table.rows:
         for cell in row.cells:
            if cell.text.strip():
               text += cell.text + "\n"
   return text

def read_resume(file_path):
   if file_path.suffix.lower() == ".pdf":
      return read_pdf(file_path)
   elif file_path.suffix.lower() == ".docx":
      return read_docx(file_path)
   else:
      return None
      
resume_folder = Path("resumes")
all_results = []

for file_path in resume_folder.iterdir():
   if file_path.suffix.lower() not in [".pdf", ".docx"]:
      continue
   print("\nProcessing:", file_path.name)
   resume_text = read_resume(file_path)
   parsed_resume = parse_resume(resume_text)
   time.sleep(5)
   result = final_score(job, parsed_resume)
   time.sleep(5)
   print("Score", result.score)
   all_results.append({
      "name": parsed_resume.name,
      "score": result.score,
      "details": result.details
   })

   all_results.sort(
      key = lambda candidate: candidate["score"],
      reverse = True
   )
   
   top_2 = all_results[:2]
   worst_2 = all_results[-2:]

   print("TOP CANDIDATES")

   for candidate in top_2:
      print(candidate["name"],
            "-",
            candidate["score"],
            "%")
      
      print(candidate["details"])

print("LOWEST 2 CANDIDATES")
for candidate in worst_2:
 print(
        candidate["name"],
        "-",
        candidate["score"],
        "%"
    )
print(candidate["details"])